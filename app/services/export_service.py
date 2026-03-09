"""导出服务：支持 TXT / DOCX / PDF 简历导出。"""

from __future__ import annotations

import re
import shutil
import subprocess
import os
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader

from app.core.config import settings
from app.services.session_store import get_version

SUPPORTED_EXPORT_FORMATS = {"txt", "docx", "pdf"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".bmp", ".gif"}


def _resolve_template(template_name: str | None) -> Path | None:
    """解析模板名称为文件路径。"""
    if not template_name:
        return None
    template_dir = settings.template_dir
    candidate = template_dir / template_name
    if not candidate.exists():
        candidate = template_dir / f"{template_name}.docx"
    if candidate.exists() and candidate.suffix.lower() == ".docx":
        return candidate
    return None


def list_templates() -> list[dict[str, str]]:
    """列出可用的 DOCX 模板。"""
    template_dir = settings.template_dir
    if not template_dir.exists():
        return []
    return [
        {"name": f.stem, "filename": f.name}
        for f in sorted(template_dir.iterdir())
        if f.suffix.lower() == ".docx" and not f.name.startswith("~$")
    ]


def _ensure_export_dir(session_id: str) -> Path:
    export_dir = settings.session_dir / session_id / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def _strip_markdown(line: str) -> str:
    cleaned = re.sub(r"^\s*[-*+]\s+", "", line.strip())
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"`(.*?)`", r"\1", cleaned)
    return cleaned.strip()


def _parse_markdown_sections(resume_text: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {}
    current = "简历内容"
    sections[current] = []

    for raw in resume_text.splitlines():
        line = raw.rstrip()
        heading_match = re.match(r"^\s*##\s*\**\s*(.+?)\s*\**\s*$", line)
        if heading_match:
            current = heading_match.group(1).strip()
            sections.setdefault(current, [])
            continue

        cleaned = _strip_markdown(line)
        if cleaned:
            sections.setdefault(current, []).append(cleaned)
    return sections


def _truncate_line(text: str, max_len: int) -> str:
    if len(text) <= max_len:
        return text
    return f"{text[: max_len - 3].rstrip()}..."


def _compact_sections(sections: dict[str, list[str]], strict: bool) -> dict[str, list[str]]:
    # 通过“按模块限额 + 行宽截断”控制一页内信息密度。
    max_lines = {
        "个人信息": 5,
        "教育背景": 4,
        "技术技能": 8,
        "实习经历": 8,
        "项目经历": 8,
    }
    if strict:
        max_lines.update({"技术技能": 6, "实习经历": 6, "项目经历": 6})

    line_len = 60 if strict else 72
    compacted: dict[str, list[str]] = {}
    for section, lines in sections.items():
        line_cap = max_lines.get(section, 5)
        trimmed = [_truncate_line(line, line_len) for line in lines[:line_cap]]
        compacted[section] = trimmed
    return compacted


def _set_page_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.font.name = "微软雅黑"
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.line_spacing = 1.15


def _add_horizontal_line(doc: Document) -> None:
    """添加一条细横线分隔符。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "999999",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc: Document, title: str) -> None:
    """添加带底部深色线的模块标题。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    # 标题下方加深蓝线
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "8",
        qn("w:space"): "1",
        qn("w:color"): "1A3C6E",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_body_paragraph(doc: Document, text: str, indent: bool = False) -> None:
    """添加正文段落（非列表项）。"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.15

    # 检测是否有加粗前缀，如 "**公司名**"
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(10.5)
        else:
            run = p.add_run(part)
            run.font.size = Pt(10.5)


def _is_entry_title(line: str) -> bool:
    """判断一行是否为经历/项目条目标题（含 · 分隔的公司·职位·时间 格式）。"""
    return bool(re.match(r'^.+[·\|].+[·\|].+$', line)) or (
        re.match(r'^\*\*.+\*\*', line) is not None and '·' in line
    )


def _add_entry_title(doc: Document, text: str) -> None:
    """添加经历/项目条目标题（如：公司 · 职位 · 时间），比正文稍大，加粗。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.2)
    cleaned = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    run = p.add_run(cleaned)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "微软雅黑"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _resolve_photo_file(photo_filename: str | None) -> Path | None:
    if not photo_filename:
        return None
    candidate = settings.upload_dir / Path(photo_filename).name
    if candidate.exists() and candidate.suffix.lower() in IMAGE_SUFFIXES:
        return candidate
    return None


def _add_header_with_photo(
    doc: Document,
    profile_lines: Iterable[str],
    photo_path: Path | None,
) -> None:
    lines = list(profile_lines)[:6]
    # 提取姓名（第一行通常是姓名）
    name = lines[0] if lines else "简历"
    info_lines = lines[1:] if len(lines) > 1 else []

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    left_cell.width = Cm(14.0)
    right_cell.width = Cm(3.0)

    # 去掉表格边框
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tblBorders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0", qn("w:color"): "auto",
        })
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # 左侧：姓名 + 联系方式
    title_p = left_cell.paragraphs[0]
    title_p.paragraph_format.space_after = Pt(2)
    name_run = title_p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.name = "微软雅黑"
    name_run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    name_run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    if info_lines:
        info_p = left_cell.add_paragraph()
        info_p.paragraph_format.space_before = Pt(2)
        info_p.paragraph_format.space_after = Pt(0)
        info_text = "  |  ".join(info_lines)
        info_run = info_p.add_run(info_text)
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 右侧：照片
    photo_paragraph = right_cell.paragraphs[0]
    photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if photo_path:
        run = photo_paragraph.add_run()
        run.add_picture(str(photo_path), width=Cm(2.8))
    else:
        # 无照片时不输出占位文字
        pass

    # 姓名下方加一条分隔线
    _add_horizontal_line(doc)


def _render_docx(resume_text: str, output_file: Path, photo_filename: str | None = None, strict: bool = False, template_path: Path | None = None) -> None:
    sections = _parse_markdown_sections(resume_text)
    sections = _compact_sections(sections, strict=strict)

    if template_path and template_path.exists():
        doc = Document(str(template_path))
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('p', 'tbl', 'sdt'):
                body.remove(child)
    else:
        doc = Document()
        _set_page_style(doc)

    profile_lines = sections.pop("个人信息", [])
    _add_header_with_photo(doc, profile_lines, _resolve_photo_file(photo_filename))

    # 经历类板块需要区分条目标题和正文
    experience_sections = {"实习经历", "项目经历", "工作经历", "科研经历", "竞赛经历"}

    for section_name, lines in sections.items():
        if not lines:
            continue

        _add_section_heading(doc, section_name)

        if section_name in experience_sections:
            for line in lines:
                if _is_entry_title(line):
                    _add_entry_title(doc, line)
                else:
                    _add_body_paragraph(doc, line, indent=True)
        else:
            for line in lines:
                _add_body_paragraph(doc, line, indent=True)

    # 内容较少时自动加大行距填充页面
    total_lines = sum(len(v) for v in sections.values()) + len(profile_lines)
    if total_lines < 20 and not strict:
        for p in doc.paragraphs:
            if p.paragraph_format.line_spacing and p.paragraph_format.line_spacing < 1.5:
                p.paragraph_format.line_spacing = 1.4
            if p.paragraph_format.space_after is not None:
                p.paragraph_format.space_after = Pt(3)

    doc.save(output_file)


def _find_soffice() -> str | None:
    env_path = os.getenv("SOFFICE_PATH", "").strip()
    if env_path:
        return env_path
    return shutil.which("soffice") or shutil.which("libreoffice")


def _docx_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    soffice = _find_soffice()
    if not soffice:
        raise RuntimeError("未找到 LibreOffice（soffice），无法导出 PDF")

    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(docx_path),
    ]
    subprocess.run(cmd, check=True, capture_output=True)
    pdf_path = out_dir / f"{docx_path.stem}.pdf"
    if not pdf_path.exists():
        raise RuntimeError("PDF 导出失败，未生成目标文件")
    return pdf_path


def _pdf_page_count(pdf_path: Path) -> int:
    reader = PdfReader(str(pdf_path))
    return len(reader.pages)


def render_docx_from_text(
    resume_text: str,
    template_name: str | None = None,
    photo_filename: str | None = None,
    compact: bool = False,
) -> bytes:
    """从文本直接渲染 DOCX 并返回字节（用于预览/下载）。"""
    template_path = _resolve_template(template_name)
    preview_dir = settings.upload_dir / "_preview"
    preview_dir.mkdir(parents=True, exist_ok=True)
    output = preview_dir / "preview.docx"
    _render_docx(resume_text, output, photo_filename=photo_filename, strict=compact, template_path=template_path)
    content = output.read_bytes()
    output.unlink(missing_ok=True)
    return content


def export_txt(session_id: str, version_id: str) -> Path:
    """导出文本并返回文件路径。"""
    version = get_version(session_id, version_id)
    export_dir = _ensure_export_dir(session_id)

    output_file = export_dir / f"resume_{version_id}.txt"
    output_file.write_text(version["resume_text"], encoding="utf-8")
    return output_file


def export_docx(session_id: str, version_id: str, photo_filename: str | None = None, template_name: str | None = None) -> Path:
    """导出一页版 DOCX（含照片位置，可选模板）。"""
    version = get_version(session_id, version_id)
    export_dir = _ensure_export_dir(session_id)
    output_file = export_dir / f"resume_{version_id}.docx"
    template_path = _resolve_template(template_name)
    _render_docx(version["resume_text"], output_file, photo_filename=photo_filename, strict=False, template_path=template_path)
    return output_file


def export_pdf(session_id: str, version_id: str, photo_filename: str | None = None, template_name: str | None = None) -> Path:
    """导出 PDF；超过一页时自动压缩一次。"""
    version = get_version(session_id, version_id)
    export_dir = _ensure_export_dir(session_id)
    template_path = _resolve_template(template_name)

    docx_file = export_dir / f"resume_{version_id}.docx"
    _render_docx(version["resume_text"], docx_file, photo_filename=photo_filename, strict=False, template_path=template_path)
    pdf_file = _docx_to_pdf(docx_file, export_dir)

    if _pdf_page_count(pdf_file) > 1:
        compact_docx = export_dir / f"resume_{version_id}_compact.docx"
        _render_docx(version["resume_text"], compact_docx, photo_filename=photo_filename, strict=True, template_path=template_path)
        compact_pdf = _docx_to_pdf(compact_docx, export_dir)
        compact_pdf.replace(pdf_file)
        compact_docx.unlink(missing_ok=True)

        if _pdf_page_count(pdf_file) > 1:
            raise RuntimeError("导出后仍超过一页，请先使用改写功能压缩内容后再导出 PDF")

    return pdf_file
